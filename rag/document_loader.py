"""
AgentOps Hub — Document Loader
================================

WHAT THIS DOES:
Reads documents from various formats (PDF, DOCX, Markdown, HTML, TXT)
and converts them into a standard format our pipeline can process.

WHY IT'S SEPARATE:
The loader doesn't know or care about chunks, embeddings, or vectors.
It has ONE job: read a file and return clean text with metadata.

REAL-WORLD ANALOGY:
In a hospital, before a doctor can read your medical records, someone
has to take the paper files, handwritten notes, and digital records
and put them all into the same electronic format. That's the loader.

DESIGN DECISIONS (why I built it this way):

1. WHY return a list of Documents (not raw strings)?
   → Because metadata matters. When the RAG system cites a source,
     it needs to know WHICH document the chunk came from, WHICH page,
     and WHAT type of document it was. Raw strings lose all that.

2. WHY a factory pattern (get loader by file extension)?
   → So adding a new file type = adding one function. You don't touch
     existing code. This is the Open/Closed Principle (SOLID).

3. WHY load ALL files from a directory at once?
   → In production, you ingest a whole knowledge base, not one file.
     Batch loading with proper error handling per file is essential.

INTERVIEW TIP:
Q: "How would you design a document ingestion pipeline?"
A: "I separate loading from processing. The loader handles format-specific
   parsing (PDF, DOCX, etc.) and returns a uniform Document object with
   text and metadata. This makes the rest of the pipeline format-agnostic.
   Adding a new format means adding one loader function — nothing else changes."
"""

import os
from pathlib import Path
from dataclasses import dataclass, field
from typing import Optional
from rich import print as rprint


# ============================================================
# Document dataclass — our universal format
# ============================================================
# DATACLASS EXPLAINED:
# A dataclass is Python's way of creating a "data container" without
# writing __init__, __repr__, etc. It's like a struct in C.
#
# WHY NOT a dictionary?
# → Dictionaries have no type checking. doc["content"] could be
#   anything. With a dataclass, doc.content is always a string.
#   Your IDE gives you autocomplete and catches typos.
# ============================================================

@dataclass
class Document:
    """
    A single document loaded from a file.
    
    This is the universal format that flows through our entire pipeline:
    Loader → Chunker → Embedder → Vector Store → Retriever
    
    Every component speaks this common language.
    """
    content: str                    # The actual text content
    metadata: dict = field(         # Information ABOUT the document
        default_factory=dict
    )
    # metadata typically contains:
    #   - source: file path or URL
    #   - file_type: "pdf", "md", "docx", etc.
    #   - title: document title (if extractable)
    #   - page: page number (for PDFs)
    #   - chunk_id: assigned later by the chunker


def load_markdown(file_path: str) -> list[Document]:
    """
    Load a Markdown file.
    
    WHY MARKDOWN FIRST: It's the simplest format — just text.
    Our sample documents are all Markdown. Perfect for development.
    In production, most internal wikis (Notion, Confluence) export as MD.
    """
    path = Path(file_path)
    
    with open(path, "r", encoding="utf-8") as f:
        content = f.read()
    
    return [Document(
        content=content,
        metadata={
            "source": str(path),
            "file_name": path.name,
            "file_type": "markdown",
            "title": _extract_md_title(content),
        }
    )]


def load_text(file_path: str) -> list[Document]:
    """Load a plain text file."""
    path = Path(file_path)
    
    with open(path, "r", encoding="utf-8") as f:
        content = f.read()
    
    return [Document(
        content=content,
        metadata={
            "source": str(path),
            "file_name": path.name,
            "file_type": "text",
        }
    )]


def load_pdf(file_path: str) -> list[Document]:
    """
    Load a PDF file using PyMuPDF.
    
    WHY PyMuPDF (not PyPDF2 or pdfplumber)?
    → PyMuPDF is the fastest PDF parser in Python (10x faster than PyPDF2).
    → It handles complex PDFs (multi-column, tables) better.
    → It's the industry standard for production RAG systems.
    
    NOTE: Returns ONE Document per PAGE. This is important because:
    1. We can track which PAGE an answer came from
    2. Pages are natural semantic boundaries
    3. If a PDF has 100 pages, we don't create one massive string
    """
    import fitz  # PyMuPDF imports as "fitz" (historical naming)
    
    path = Path(file_path)
    documents = []
    
    pdf = fitz.open(str(path))
    for page_num, page in enumerate(pdf, start=1):
        text = page.get_text()
        if text.strip():  # Skip empty pages
            documents.append(Document(
                content=text,
                metadata={
                    "source": str(path),
                    "file_name": path.name,
                    "file_type": "pdf",
                    "page": page_num,
                    "total_pages": len(pdf),
                }
            ))
    pdf.close()
    
    return documents


def load_docx(file_path: str) -> list[Document]:
    """
    Load a Word document (.docx).
    
    WHY: Many company docs are in Word format — policies, procedures,
    contracts, reports. This is non-negotiable for enterprise RAG.
    """
    from docx import Document as DocxDocument
    
    path = Path(file_path)
    doc = DocxDocument(str(path))
    
    # Extract all paragraph text
    paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
    content = "\n\n".join(paragraphs)
    
    return [Document(
        content=content,
        metadata={
            "source": str(path),
            "file_name": path.name,
            "file_type": "docx",
        }
    )]


def load_html(file_path: str) -> list[Document]:
    """
    Load an HTML file, stripping tags to get clean text.
    
    WHY: Internal wikis, saved web pages, exported Confluence pages.
    """
    from bs4 import BeautifulSoup
    
    path = Path(file_path)
    
    with open(path, "r", encoding="utf-8") as f:
        soup = BeautifulSoup(f.read(), "html.parser")
    
    # Remove script and style elements
    for tag in soup(["script", "style"]):
        tag.decompose()
    
    content = soup.get_text(separator="\n")
    # Clean up excessive whitespace
    lines = [line.strip() for line in content.splitlines() if line.strip()]
    content = "\n".join(lines)
    
    title = soup.title.string if soup.title else None
    
    return [Document(
        content=content,
        metadata={
            "source": str(path),
            "file_name": path.name,
            "file_type": "html",
            "title": title,
        }
    )]


# ============================================================
# Loader registry — maps file extensions to loader functions
# ============================================================
# This is a simple FACTORY PATTERN:
# Instead of a giant if/elif chain, we map extensions to functions.
# Adding a new format = adding one entry to this dictionary.
#
# INTERVIEW TIP:
# Q: "How would you add support for a new file format?"
# A: "I'd write a load_newformat() function and add it to the
#    LOADER_REGISTRY. Nothing else in the codebase changes."
# ============================================================

LOADER_REGISTRY: dict[str, callable] = {
    ".md": load_markdown,
    ".markdown": load_markdown,
    ".txt": load_text,
    ".pdf": load_pdf,
    ".docx": load_docx,
    ".html": load_html,
    ".htm": load_html,
}


def load_document(file_path: str) -> list[Document]:
    """
    Load a single document. Automatically detects format by extension.
    
    Args:
        file_path: Path to the document file
        
    Returns:
        List of Document objects (PDFs return one per page, others return one)
        
    Raises:
        ValueError: If file format is not supported
        FileNotFoundError: If file doesn't exist
    """
    path = Path(file_path)
    
    if not path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")
    
    extension = path.suffix.lower()
    
    if extension not in LOADER_REGISTRY:
        raise ValueError(
            f"Unsupported file format: {extension}. "
            f"Supported: {list(LOADER_REGISTRY.keys())}"
        )
    
    loader = LOADER_REGISTRY[extension]
    return loader(file_path)


def load_directory(directory_path: str) -> list[Document]:
    """
    Load ALL supported documents from a directory.
    
    WHY THIS EXISTS:
    In production, you don't ingest files one by one. You point the
    system at a folder (or S3 bucket, or Confluence space) and say "ingest everything."
    
    IMPORTANT: This handles errors per-file. If one PDF is corrupted,
    it logs the error and continues with the rest. In production,
    you NEVER let one bad file crash the entire ingestion.
    """
    directory = Path(directory_path)
    
    if not directory.exists():
        raise FileNotFoundError(f"Directory not found: {directory_path}")
    
    all_documents = []
    supported_extensions = set(LOADER_REGISTRY.keys())
    
    # Find all supported files (including subdirectories)
    files = [
        f for f in directory.rglob("*")
        if f.is_file() and f.suffix.lower() in supported_extensions
    ]
    
    if not files:
        rprint(f"[yellow]⚠️  No supported documents found in {directory_path}[/yellow]")
        return []
    
    rprint(f"[cyan]📁 Found {len(files)} documents to load[/cyan]")
    
    for file_path in sorted(files):
        try:
            docs = load_document(str(file_path))
            all_documents.extend(docs)
            rprint(f"  [green]✅ Loaded: {file_path.name} ({len(docs)} doc(s))[/green]")
        except Exception as e:
            # Log error but continue — don't let one bad file crash everything
            rprint(f"  [red]❌ Failed: {file_path.name} — {e}[/red]")
    
    rprint(f"[cyan]📄 Total documents loaded: {len(all_documents)}[/cyan]")
    return all_documents


# ============================================================
# Helper functions
# ============================================================

def _extract_md_title(content: str) -> Optional[str]:
    """Extract the first H1 heading from Markdown content as the title."""
    for line in content.split("\n"):
        line = line.strip()
        if line.startswith("# ") and not line.startswith("##"):
            return line[2:].strip()
    return None


# ============================================================
# Self-test: Run this file directly to verify it works
# ============================================================
if __name__ == "__main__":
    import sys
    
    # Default to our sample documents directory
    doc_dir = sys.argv[1] if len(sys.argv) > 1 else "rag/documents"
    
    rprint(f"\n[bold]🔍 Testing Document Loader[/bold]")
    rprint(f"   Directory: {doc_dir}\n")
    
    docs = load_directory(doc_dir)
    
    if docs:
        rprint(f"\n[bold]📋 Sample output (first document):[/bold]")
        rprint(f"   Title:    {docs[0].metadata.get('title', 'N/A')}")
        rprint(f"   Source:   {docs[0].metadata.get('file_name', 'N/A')}")
        rprint(f"   Type:     {docs[0].metadata.get('file_type', 'N/A')}")
        rprint(f"   Content:  {docs[0].content[:200]}...")